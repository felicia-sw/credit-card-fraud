"""
Generates the research paper (JUISI template) as a Word .docx.

Reproducible artifact: run `python create_paper.py` from the repo root. It embeds the
figures from outputs/ and the (deterministic, SEED=42) results of 02_analysis.ipynb.

Output: paper/Deteksi_Fraud_KNN_vs_LR_JUISI.docx
"""
import os
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = 'paper'
FIG = 'outputs'
os.makedirs(OUT_DIR, exist_ok=True)

doc = Document()

# ── Page setup: 192 x 262 mm trim, single column ──
sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(192), Mm(262)
sec.top_margin = sec.bottom_margin = Mm(20)
sec.left_margin = sec.right_margin = Mm(18)

normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(10)

# Running header
hdr = sec.header.paragraphs[0]
hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
_r = hdr.add_run('Volume xx Nomor yy\ne-ISSN: 2477-5452; p-ISSN: 2460-1306')
_r.font.name = 'Times New Roman'; _r.font.size = Pt(8); _r.italic = True

_tbl = [0]; _fig = [0]; _eq = [0]


def _font(run, size=10, bold=False, italic=False, sup=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if sup:
        run.font.superscript = True


def para(text='', size=10, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         space_after=6, space_before=0, indent=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if indent:
        p.paragraph_format.first_line_indent = Mm(7)
    if text:
        _font(p.add_run(text), size, bold, italic)
    return p


def rich(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6, indent=False):
    """parts: list of (text, dict-of-font-kwargs)."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Mm(7)
    for text, kw in parts:
        _font(p.add_run(text), **kw)
    return p


def heading(text):
    return para(text, size=10, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=10, space_after=4)


def subheading(text):
    return para(text, size=10, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=6, space_after=3)


def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    _font(p.add_run(text), 10)
    return p


def equation(expr):
    _eq[0] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    _font(p.add_run(expr), 10, italic=True)
    _font(p.add_run(f'      ({_eq[0]})'), 10)
    return p


def _cell_border(cell, edges):
    tcPr = cell._tc.get_or_add_tcPr()
    tb = tcPr.find(qn('w:tcBorders'))
    if tb is None:
        tb = OxmlElement('w:tcBorders'); tcPr.append(tb)
    for edge in edges:
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '8')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), '000000')
        tb.append(e)


def add_table(caption, headers, rows):
    _tbl[0] += 1
    para(f'Tabel {_tbl[0]}. {caption}', size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=8, space_after=2)
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        _font(c.paragraphs[0].add_run(h), 9, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            _font(cells[i].paragraphs[0].add_run(str(val)), 9)
    for c in t.rows[0].cells:
        _cell_border(c, ['top', 'bottom'])
    for c in t.rows[-1].cells:
        _cell_border(c, ['bottom'])
    para('', space_after=6)
    return t


def add_figure(path, caption, width_mm=135):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    if os.path.exists(path):
        p.add_run().add_picture(path, width=Mm(width_mm))
    else:
        _font(p.add_run(f'[gambar tidak ditemukan: {path}]'), 9, italic=True)
    _fig[0] += 1
    para(f'Gambar {_fig[0]}. {caption}', size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)


# ══════════════════════════════ TITLE BLOCK ══════════════════════════════
para('Perbandingan Algoritma K-Nearest Neighbor dan Logistic Regression untuk '
     'Deteksi Penipuan Kartu Kredit pada Data Tidak Seimbang dengan Penerapan SMOTE',
     size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para('Comparison of K-Nearest Neighbor and Logistic Regression Algorithms for '
     'Credit Card Fraud Detection on Imbalanced Data Using SMOTE',
     size=16, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

rich([('Felicia Sword', dict(size=14)), ('1*', dict(size=14, sup=True))],
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
rich([('1', dict(size=12, italic=True, sup=True)),
      ('[Program Studi …, Nama Universitas, Kota, Indonesia] — Tim Nekat part 100',
       dict(size=12, italic=True))],
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
rich([('*', dict(size=10, italic=True, sup=True)),
      ('Korespondensi: [alamat email] (NIM 0706012410012)', dict(size=10, italic=True))],
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

# ── Abstract (ID) ──
rich([('Abstrak: ', dict(size=10, bold=True)),
      ('Penipuan kartu kredit merupakan ancaman serius dalam ekosistem keuangan digital, dengan '
       'karakteristik utama berupa ketidakseimbangan kelas yang ekstrem (proporsi fraud kurang dari 1%). '
       'Penelitian ini membandingkan dua algoritma klasifikasi standar — K-Nearest Neighbor (KNN) dan '
       'Logistic Regression (LR) — untuk mendeteksi transaksi penipuan pada dataset Sparkov (1,85 juta '
       'transaksi; subsampel terstratifikasi 100.000 baris; rasio fraud 0,52%), serta menganalisis '
       'pengaruh penerapan Synthetic Minority Over-sampling Technique (SMOTE). Praproses meliputi rekayasa '
       'fitur temporal, jarak haversine, dan fitur perilaku per-kartu yang dihitung tanpa kebocoran data '
       '(leak-free), one-hot dan label encoding, serta standardisasi z-score. Model dievaluasi dengan '
       'pembagian terstratifikasi 80:20 menggunakan confusion matrix, accuracy, precision, recall, '
       'F1-score, ROC-AUC, dan PR-AUC, dilengkapi penyetelan ambang keputusan (threshold tuning), selang '
       'kepercayaan dari lima pengulangan, serta uji signifikansi McNemar. Hasil menunjukkan KNN+SMOTE '
       'memperoleh F1 tertinggi (0,537) dan KNN baseline memperoleh precision serta PR-AUC tertinggi, '
       'sedangkan LR+SMOTE unggul pada recall (0,769) dan ROC-AUC (0,901) namun memiliki PR-AUC terendah '
       '(0,188). Uji McNemar mengonfirmasi keunggulan KNN atas LR secara signifikan (p < 0,01). Temuan '
       'utama: pada data sangat tidak seimbang ROC-AUC dapat menyesatkan sehingga PR-AUC lebih informatif; '
       'KNN+SMOTE direkomendasikan sebagai detektor yang seimbang, sedangkan LR+SMOTE sesuai bila recall '
       'menjadi prioritas utama.', dict(size=10))], space_after=4)
rich([('Kata Kunci: ', dict(size=10, bold=True)),
      ('deteksi fraud; kartu kredit; K-Nearest Neighbor; regresi logistik; SMOTE; ketidakseimbangan kelas; PR-AUC',
       dict(size=10, italic=True))], space_after=6)

# ── Abstract (EN) ──
rich([('Abstract: ', dict(size=10, bold=True)),
      ('Credit card fraud is a serious threat in the digital financial ecosystem, characterised mainly by '
       'extreme class imbalance (fraud accounts for less than 1% of transactions). This study compares two '
       'standard classification algorithms — K-Nearest Neighbor (KNN) and Logistic Regression (LR) — for '
       'detecting fraudulent transactions on the Sparkov dataset (1.85 million transactions; a stratified '
       'subsample of 100,000 rows; 0.52% fraud rate), and analyses the effect of the Synthetic Minority '
       'Over-sampling Technique (SMOTE). Preprocessing includes temporal feature engineering, haversine '
       'distance, leak-free per-card behavioural features, one-hot and label encoding, and z-score '
       'standardisation. Models are evaluated with a stratified 80:20 split using the confusion matrix, '
       'accuracy, precision, recall, F1-score, ROC-AUC, and PR-AUC, complemented by decision-threshold '
       'tuning, confidence intervals over five repetitions, and the McNemar significance test. KNN+SMOTE '
       'achieves the highest F1 (0.537) and KNN baseline the highest precision and PR-AUC, whereas LR+SMOTE '
       'leads on recall (0.769) and ROC-AUC (0.901) but has the lowest PR-AUC (0.188). McNemar confirms '
       'that KNN significantly outperforms LR (p < 0.01). The main finding: under severe imbalance ROC-AUC '
       'can be misleading and PR-AUC is more informative; KNN+SMOTE is recommended as a balanced detector, '
       'while LR+SMOTE suits cases where recall is the priority.', dict(size=10))], space_after=4)
rich([('Keywords: ', dict(size=10, bold=True)),
      ('fraud detection; credit card; K-Nearest Neighbor; logistic regression; SMOTE; class imbalance; PR-AUC',
       dict(size=10, italic=True))], space_after=6)

para('Naskah diterima [tanggal]; direvisi [tanggal]; dipublikasi [tanggal].',
     size=8, italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)

# ══════════════════════════════ 1. PENDAHULUAN ══════════════════════════════
heading('1.  Pendahuluan')
para('Penipuan kartu kredit (credit card fraud) merupakan salah satu ancaman paling signifikan dalam '
     'ekosistem keuangan digital global. Berdasarkan laporan Federal Trade Commission (FTC), kerugian '
     'individu akibat penipuan finansial di Amerika Serikat pada tahun 2022 mencapai lebih dari USD 8,8 '
     'miliar, meningkat 30% dibandingkan tahun sebelumnya (Federal Trade Commission, 2023). Di Indonesia, '
     'Otoritas Jasa Keuangan (OJK) mencatat kerugian masyarakat akibat penipuan finansial digital mencapai '
     'Rp 7,8 triliun sepanjang November 2024, sementara kejahatan siber melonjak 550 persen pada periode '
     'yang sama (OJK, 2026). Seiring pesatnya pertumbuhan layanan finansial digital, deteksi penipuan '
     'menjadi komponen kritis dalam sistem keamanan finansial modern.', indent=True)
para('Tantangan utama dalam deteksi fraud adalah ketidakseimbangan kelas yang parah (severe class '
     'imbalance): proporsi transaksi penipuan pada data nyata umumnya kurang dari 1%, sehingga model '
     'klasifikasi cenderung bias terhadap kelas mayoritas (transaksi sah). Akibatnya, metrik accuracy '
     'menjadi menyesatkan — sebuah model yang selalu memprediksi “bukan fraud” dapat mencapai akurasi '
     '>99% namun gagal total mendeteksi penipuan. Penanganan ketidakseimbangan, misalnya melalui '
     'oversampling sintetis SMOTE (Chawla et al., 2002), serta pemilihan metrik evaluasi yang tepat '
     'menjadi sangat penting.', indent=True)
para('Terdapat beragam algoritma klasifikasi yang dapat digunakan, masing-masing dengan asumsi, '
     'kelebihan, dan keterbatasan berbeda. Studi yang melakukan analisis perbandingan sistematis antara '
     'algoritma klasik pada satu dataset dengan protokol evaluasi yang konsisten — termasuk penanganan '
     'ketidakseimbangan melalui SMOTE — masih terbatas dalam literatur, padahal pemilihan algoritma yang '
     'tepat memiliki implikasi praktis signifikan bagi praktisi fintech.', indent=True)
para('Dari kondisi tersebut, dirumuskan permasalahan utama: di antara dua algoritma klasifikasi standar '
     '(K-Nearest Neighbor dan Logistic Regression), algoritma manakah yang memberikan performa terbaik '
     'untuk mengklasifikasikan transaksi credit card fraud pada dataset yang sangat tidak seimbang, dan '
     'bagaimana penerapan SMOTE memengaruhi performa relatifnya? Kontribusi penelitian ini adalah: '
     '(1) perbandingan empiris KNN vs LR dengan protokol evaluasi konsisten; (2) analisis pengaruh SMOTE '
     'pada masing-masing algoritma; dan (3) penekanan pada PR-AUC, threshold tuning, serta uji '
     'signifikansi statistik untuk memperkuat keandalan kesimpulan. Sisa artikel disusun sebagai berikut: '
     'Bagian 2 membahas kajian pustaka, Bagian 3 metode, Bagian 4 hasil dan pembahasan, dan Bagian 5 '
     'kesimpulan dan saran.', indent=True)

# ══════════════════════════════ 2. KAJIAN PUSTAKA ══════════════════════════════
heading('2.  Kajian Pustaka')
para('Penelitian deteksi fraud kerap terkendala ketersediaan data nyata karena alasan kerahasiaan. '
     'Untuk mengatasinya, sejumlah simulator dan dataset sintetis dikembangkan, antara lain PaySim yang '
     'mensimulasikan transaksi uang elektronik untuk riset fraud (Lopez-Rojas et al., 2016) dan Sparkov '
     'Data Generation Tool yang menghasilkan dataset transaksi kartu kredit yang digunakan pada penelitian '
     'ini (Harris, 2016; Kaggle/Shenoy, 2020). Grover et al. (2022) menyusun Fraud Dataset Benchmark yang '
     'menstandarkan evaluasi lintas dataset fraud dan menekankan pentingnya protokol yang konsisten.', indent=True)
para('K-Nearest Neighbor (KNN) adalah algoritma berbasis instans yang mengklasifikasikan sampel '
     'berdasarkan mayoritas kelas dari k tetangga terdekat dalam ruang fitur; karena berbasis jarak, KNN '
     'sangat sensitif terhadap skala fitur sehingga standardisasi menjadi wajib. Logistic Regression (LR) '
     'adalah model linier yang memodelkan probabilitas kelas melalui fungsi logistik dan menghasilkan '
     'estimasi probabilitas yang terkalibrasi serta mudah diinterpretasikan. Kedua algoritma ini '
     'merupakan baseline klasik yang lazim dijadikan pembanding.', indent=True)
para('Untuk menangani ketidakseimbangan kelas, SMOTE membangkitkan sampel minoritas sintetis melalui '
     'interpolasi antar tetangga terdekat kelas minoritas (Chawla et al., 2002), dan harus diterapkan '
     'hanya pada data latih agar evaluasi tetap realistis. Terkait evaluasi, Saito dan Rehmsmeier (2015) '
     'menunjukkan bahwa pada data yang sangat tidak seimbang kurva Precision–Recall (dan PR-AUC) lebih '
     'informatif daripada kurva ROC, karena ROC dapat terlihat optimistis meskipun precision rendah. '
     'Temuan tersebut menjadi dasar penggunaan PR-AUC sebagai metrik pelengkap dalam penelitian ini. '
     'Implementasi seluruh model menggunakan pustaka scikit-learn (Pedregosa et al., 2011).', indent=True)

# ══════════════════════════════ 3. METODE ══════════════════════════════
heading('3.  Metode')
subheading('3.1  Dataset')
para('Penelitian menggunakan Credit Card Transactions Fraud Detection Dataset yang dibangkitkan dengan '
     'Sparkov (Kaggle/Shenoy, 2020; Harris, 2016), berisi 1.852.394 transaksi pada periode 2019–2020 '
     'dengan rasio fraud sangat rendah. Untuk efisiensi komputasi, diambil subsampel terstratifikasi '
     'sebanyak 100.000 transaksi yang mempertahankan rasio fraud asli (0,52%; 521 transaksi fraud). '
     'Variabel penelitian dirangkum pada Tabel 1.', indent=True)
add_table('Variabel penelitian',
          ['Jenis', 'Nama', 'Keterangan'],
          [['Dependen (Y)', 'is_fraud', 'Biner: 1 = fraud, 0 = sah'],
           ['Independen X1', 'category', 'Kategori merchant (14 kategori, one-hot)'],
           ['Independen X2', 'gender', 'Jenis kelamin pemegang kartu (label encoding)'],
           ['Independen X3', 'age', 'Usia pemegang kartu (diturunkan dari dob)'],
           ['Independen X4', 'city_pop', 'Populasi kota pemegang kartu'],
           ['Independen X5', 'hour', 'Jam transaksi (0–23)'],
           ['Independen X6', 'day_of_week', 'Hari transaksi (one-hot)'],
           ['Independen X7', 'is_weekend', 'Indikator akhir pekan (1/0)'],
           ['Independen X8', 'distance', 'Jarak haversine kartu–merchant (km)'],
           ['Independen X9', 'amt', 'Nilai transaksi (USD) — perluasan'],
           ['Perilaku', 'amt_zscore_card', 'Nilai transaksi relatif terhadap riwayat kartu'],
           ['Perilaku', 'hours_since_prev', 'Jam sejak transaksi sebelumnya kartu'],
           ['Perilaku', 'txns_24h', 'Jumlah transaksi kartu dalam 24 jam terakhir']])

subheading('3.2  Praproses dan Rekayasa Fitur')
para('Fitur temporal (hour, day_of_week, is_weekend) diekstraksi dari stempel waktu transaksi, dan usia '
     '(age) dihitung dari tanggal lahir. Jarak geografis antara pemegang kartu dan merchant dihitung '
     'menggunakan formula haversine pada Persamaan (1), dengan φ lintang, λ bujur (radian), dan r = 6.371 '
     'km jari-jari Bumi.', indent=True)
equation('d = 2r · arcsin( √( sin²(Δφ/2) + cosφ₁ · cosφ₂ · sin²(Δλ/2) ) )')
para('Selain itu dibangun tiga fitur perilaku per-kartu yang menangkap pola penyalahgunaan: '
     'amt_zscore_card (seberapa tidak wajar nilai transaksi dibanding riwayat kartu), hours_since_prev, '
     'dan txns_24h. Fitur-fitur ini dihitung pada keseluruhan 1,85 juta transaksi sebelum subsampling — '
     'karena membutuhkan riwayat penuh tiap kartu — dan bersifat bebas kebocoran (leak-free): setiap '
     'transaksi hanya menggunakan riwayatnya sendiri di masa lalu. Sebagai contoh, amt_zscore_card '
     'dihitung dengan Persamaan (2) menggunakan rata-rata (μ) dan simpangan baku (σ) transaksi kartu '
     'tersebut sebelum transaksi berjalan.', indent=True)
equation('z = (amt − μ_riwayat) / σ_riwayat')
para('Variabel kategorik category dan day_of_week dikodekan dengan one-hot encoding (drop-first untuk '
     'menghindari multikolinearitas), gender dengan label encoding, dan seluruh fitur numerik '
     'distandardisasi menggunakan z-score (StandardScaler). Standardisasi sangat krusial bagi KNN yang '
     'berbasis jarak.', indent=True)

subheading('3.3  Penanganan Ketidakseimbangan Kelas')
para('Dua kondisi dibandingkan: (a) baseline tanpa resampling, dan (b) SMOTE. SMOTE diterapkan hanya '
     'pada data latih untuk membangkitkan sampel kelas minoritas sintetis; data uji tidak dimodifikasi '
     'agar evaluasi tetap mencerminkan distribusi nyata (Chawla et al., 2002).', indent=True)

subheading('3.4  Model dan Pemilihan Hyperparameter')
para('Dua algoritma dibandingkan, masing-masing pada kondisi baseline dan SMOTE, sehingga terdapat empat '
     'varian model: KNN Baseline, KNN+SMOTE, LR Baseline, dan LR+SMOTE. Nilai k optimal pada KNN dipilih '
     'melalui 5-fold stratified cross-validation pada data latih dengan kriteria F1 kelas fraud, dari '
     'kandidat k ∈ {1, 3, 5, 7, 9, 11, 15, 21}. Logistic Regression dilatih dengan maksimum 1.000 iterasi. '
     'Seluruh proses menggunakan random seed tetap (42) sehingga hasil bersifat reproducible.', indent=True)

subheading('3.5  Skema Evaluasi')
para('Data dibagi secara terstratifikasi 80:20 menjadi data latih dan uji. Metrik yang dilaporkan '
     'meliputi confusion matrix, accuracy, precision, recall, F1-score (Persamaan 3), ROC-AUC, dan PR-AUC. '
     'Perbandingan utama difokuskan pada F1 dan recall kelas fraud (minoritas).', indent=True)
equation('F1 = 2 · (precision · recall) / (precision + recall)')
para('Sebagai analisis lanjutan, dilakukan penyetelan ambang keputusan (threshold tuning) dengan memilih '
     'ambang yang memaksimalkan F1 dari kurva precision–recall, evaluasi stabilitas melalui selang '
     'kepercayaan (rata-rata ± simpangan baku) atas lima pembagian acak berbeda, serta uji signifikansi '
     'McNemar (Persamaan 4) untuk menguji apakah perbedaan kebenaran prediksi antar algoritma bersifat '
     'signifikan secara statistik (b dan c adalah jumlah prediksi yang berbeda hasil benarnya antar dua '
     'model).', indent=True)
equation('χ² = (|b − c| − 1)² / (b + c)')

# ══════════════════════════════ 4. HASIL DAN PEMBAHASAN ══════════════════════════════
heading('4.  Hasil dan Pembahasan')

subheading('4.1  Analisis Data Eksploratif (EDA)')
para('Gambar 1 menegaskan ketidakseimbangan kelas yang ekstrem: hanya 521 dari 100.000 transaksi (0,52%) '
     'berlabel fraud. Hal ini menggarisbawahi bahwa accuracy bukan metrik yang memadai dan F1/recall kelas '
     'fraud menjadi acuan utama.', indent=True)
add_figure(f'{FIG}/eda_class_distribution.png', 'Distribusi kelas transaksi (sah vs fraud).')
para('Tingkat fraud sangat bervariasi antar kategori merchant (Gambar 2), dengan kategori belanja daring '
     'dan beberapa kategori lain memiliki tingkat fraud jauh lebih tinggi. Analisis fitur perilaku '
     'menunjukkan sinyal terkuat: rata-rata amt_zscore_card pada transaksi fraud mencapai ≈3,82 (yakni '
     'sekitar 3,8 simpangan baku di atas kebiasaan belanja kartu), dibandingkan ≈0,00 pada transaksi sah. '
     'Sebaliknya, fitur kecepatan (hours_since_prev, txns_24h) hanya memisahkan kelas secara lemah '
     '(|korelasi| < 0,02) — sebuah temuan negatif yang tetap informatif. Matriks korelasi pada Gambar 3 '
     'mengonfirmasi bahwa amt dan amt_zscore_card adalah prediktor numerik terkuat.', indent=True)
add_figure(f'{FIG}/eda_fraud_by_category.png', 'Tingkat fraud per kategori merchant.')
add_figure(f'{FIG}/correlation_heatmap.png', 'Matriks korelasi Pearson antar fitur numerik.')

subheading('4.2  Pemilihan Nilai k pada KNN')
para('Cross-validation memilih k = 1 baik pada kondisi baseline maupun SMOTE, karena nilai F1 menurun '
     'monoton seiring bertambahnya k. Hal ini merupakan gejala ketidakseimbangan: pada k > 1, lingkungan '
     'tetangga didominasi transaksi sah sehingga kelas minoritas nyaris tidak pernah memenangkan suara '
     'mayoritas. Implikasinya, probabilitas KNN bersifat biner (0/1) sehingga sulit disetel ambangnya — '
     'isu yang dibahas pada Subbagian 4.5.', indent=True)

subheading('4.3  Perbandingan Kinerja Model')
para('Tabel 2 menyajikan kinerja keempat model pada ambang default 0,5. KNN+SMOTE memperoleh F1 tertinggi '
     '(0,537) dan KNN baseline precision tertinggi (0,644) serta PR-AUC tertinggi (0,294). Sebaliknya, '
     'LR+SMOTE memperoleh recall tertinggi (0,769) dan ROC-AUC tertinggi (0,901), tetapi PR-AUC terendah '
     '(0,188) akibat precision yang sangat rendah (0,040). LR baseline praktis gagal mendeteksi fraud pada '
     'ambang default (recall 0,029) — ilustrasi klasik “accuracy paradox”.', indent=True)
add_table('Perbandingan kinerja empat model (ambang default 0,5)',
          ['Model', 'Accuracy', 'Precision', 'Recall', 'F1', 'ROC-AUC', 'PR-AUC'],
          [['KNN Baseline', '0,9958', '0,6438', '0,4519', '0,5311', '0,7253', '0,2938'],
           ['KNN + SMOTE', '0,9950', '0,5179', '0,5577', '0,5370', '0,7775', '0,2911'],
           ['LR Baseline', '0,9943', '0,1875', '0,0288', '0,0500', '0,8369', '0,2676'],
           ['LR + SMOTE', '0,9016', '0,0395', '0,7692', '0,0752', '0,9007', '0,1884']])
para('Kurva ROC (Gambar 4) memperlihatkan keunggulan kemampuan pemeringkatan LR, namun kurva '
     'precision–recall (Gambar 5) — yang lebih relevan pada data tidak seimbang — menunjukkan KNN '
     'lebih unggul. Matriks konfusi keempat model disajikan pada Gambar 6.', indent=True)
add_figure(f'{FIG}/roc_curves_comparison.png', 'Kurva ROC keempat model.')
add_figure(f'{FIG}/precision_recall_curves.png', 'Kurva Precision–Recall keempat model.')
add_figure(f'{FIG}/confusion_matrices.png', 'Matriks konfusi keempat model.')

subheading('4.4  Penyetelan Ambang Keputusan (Threshold Tuning)')
para('Tabel 3 menunjukkan operating point yang memaksimalkan F1. Untuk LR baseline, penyetelan ambang ke '
     '≈0,085 menaikkan F1 dari 0,050 menjadi 0,490 — menegaskan bahwa kegagalan pada ambang default murni '
     'merupakan persoalan ambang, bukan model yang rusak. Sementara itu KNN tidak berubah karena pada '
     'k = 1 probabilitasnya bersifat biner.', indent=True)
add_table('Hasil penyetelan ambang (maksimisasi F1)',
          ['Model', 'Ambang', 'Precision', 'Recall', 'F1'],
          [['KNN Baseline', '1,000', '0,644', '0,452', '0,531'],
           ['KNN + SMOTE', '1,000', '0,518', '0,558', '0,537'],
           ['LR Baseline', '0,085', '0,534', '0,452', '0,490'],
           ['LR + SMOTE', '0,970', '0,343', '0,452', '0,390']])

subheading('4.5  KNN dengan Pembobotan Jarak')
para('Untuk mengatasi probabilitas biner pada k = 1, diuji KNN dengan pembobotan jarak '
     '(weights=distance). Ketika k dipilih ulang melalui CV, hasilnya tetap k = 1 dan identik dengan '
     'KNN seragam (bobot satu tetangga tidak relevan). Namun pada k yang lebih besar (k = 21), pembobotan '
     'jarak menghasilkan probabilitas kontinu (840 nilai unik vs 22 pada bobot seragam) sehingga model '
     'menjadi dapat disetel ambangnya; setelah penyetelan, F1 mencapai ≈0,570 — hasil KNN terbaik dalam '
     'studi ini, melampaui seluruh model pada ambang default.', indent=True)

subheading('4.6  Stabilitas dan Uji Signifikansi')
para('Selang kepercayaan atas lima pembagian acak (Tabel 4) menunjukkan hasil yang stabil: F1 KNN+SMOTE '
     '0,517 ± 0,027 dan recall LR+SMOTE 0,754 ± 0,013. Karena selangnya sempit, perbedaan antar model '
     'bersifat nyata. Uji McNemar (Tabel 5) mengonfirmasi bahwa KNN secara signifikan lebih unggul '
     'dibanding LR baik pada kondisi baseline (p = 0,001) maupun SMOTE (p < 0,0001).', indent=True)
add_table('Selang kepercayaan (rata-rata ± simpangan baku, 5 pengulangan)',
          ['Model', 'F1', 'Recall', 'ROC-AUC', 'PR-AUC'],
          [['KNN Baseline', '0,512 ± 0,054', '0,440 ± 0,053', '0,719 ± 0,026', '0,275 ± 0,057'],
           ['KNN + SMOTE', '0,517 ± 0,027', '0,544 ± 0,036', '0,771 ± 0,018', '0,271 ± 0,029'],
           ['LR Baseline', '0,068 ± 0,027', '0,038 ± 0,016', '0,817 ± 0,016', '0,268 ± 0,033'],
           ['LR + SMOTE', '0,072 ± 0,002', '0,754 ± 0,013', '0,896 ± 0,008', '0,172 ± 0,013']])
add_table('Uji signifikansi McNemar (A vs B, prediksi pada data uji)',
          ['Perbandingan (A vs B)', 'χ²', 'p-value', 'Kesimpulan'],
          [['LR Baseline vs KNN Baseline', '10,588', '0,0011', 'Signifikan (KNN unggul)'],
           ['LR + SMOTE vs KNN + SMOTE', '1764,013', '< 0,0001', 'Signifikan (KNN unggul)']])

subheading('4.7  Pembahasan')
para('Tiga temuan utama menonjol. Pertama, pemilihan metrik menentukan kesimpulan: ROC-AUC menempatkan '
     'LR+SMOTE sebagai yang terbaik (0,901), tetapi PR-AUC — yang memperhitungkan precision pada data '
     'tidak seimbang — justru mengunggulkan KNN. Recall tinggi LR+SMOTE (0,769) ditebus dengan banjir '
     'false positive (precision 0,040), sejalan dengan Saito dan Rehmsmeier (2015). Kedua, pengaruh SMOTE '
     'bergantung pada algoritma: SMOTE membantu KNN (F1 0,531→0,537; recall naik) karena fitur perilaku '
     'memberi struktur minoritas yang bermakna untuk diinterpolasi, sedangkan bagi LR, SMOTE menaikkan '
     'recall secara drastis namun menghancurkan precision — dan penyetelan ambang pada LR baseline '
     'mencapai keseimbangan yang lebih sehat. Ketiga, fitur perilaku amt_zscore_card terbukti menjadi '
     'pendorong utama peningkatan kinerja. Secara keseluruhan, untuk detektor yang dapat diterapkan dan '
     'seimbang, KNN+SMOTE direkomendasikan, sedangkan LR+SMOTE sesuai bila menangkap fraud sebanyak '
     'mungkin (recall) merupakan prioritas dengan toleransi terhadap false positive.', indent=True)

# ══════════════════════════════ 5. KESIMPULAN DAN SARAN ══════════════════════════════
heading('5.  Kesimpulan dan Saran')
para('Penelitian ini membandingkan KNN dan Logistic Regression untuk deteksi penipuan kartu kredit pada '
     'dataset Sparkov yang sangat tidak seimbang, dengan dan tanpa SMOTE. Pada metrik yang seimbang '
     '(F1, precision, PR-AUC), KNN — khususnya KNN+SMOTE (F1 0,537) — terbukti lebih unggul dan keunggulan '
     'tersebut signifikan secara statistik menurut uji McNemar (p < 0,01). Logistic Regression dengan '
     'SMOTE unggul pada recall (0,769) dan ROC-AUC (0,901), namun precision-nya sangat rendah sehingga '
     'PR-AUC-nya terendah. SMOTE terbukti membantu KNN tetapi bagi LR lebih efektif digantikan oleh '
     'penyetelan ambang. Penelitian menegaskan pentingnya melaporkan PR-AUC, bukan hanya ROC-AUC, pada '
     'data tidak seimbang.', indent=True)
para('Saran untuk penelitian lanjutan: (1) menambah rekayasa fitur perilaku lain dan menggunakan sampel '
     'lebih besar hingga dataset penuh; (2) menguji algoritma berbasis ensembel pohon (mis. Random Forest '
     'atau Gradient Boosting) sebagai pembanding; (3) menerapkan kalibrasi probabilitas atau pembobotan '
     'jarak agar KNN dapat disetel ambangnya secara wajar; dan (4) mengeksplorasi ambang sensitif-biaya '
     'sesuai toleransi false positive lembaga keuangan.', indent=True)

# ══════════════════════════════ UCAPAN TERIMA KASIH ══════════════════════════════
heading('Ucapan Terima Kasih')
para('Penulis mengucapkan terima kasih kepada dosen pengampu mata kuliah Statistika dan Probabilitas '
     'atas bimbingannya, serta kepada penyedia dataset Sparkov/Kaggle.', indent=True)

# ══════════════════════════════ DAFTAR PUSTAKA ══════════════════════════════
heading('Daftar Pustaka')
refs = [
    'Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: Synthetic minority '
    'over-sampling technique. Journal of Artificial Intelligence Research, 16, 321–357.',
    'Federal Trade Commission. (2023, Februari 21). New FTC data show consumers reported losing nearly '
    '$8.8 billion to scams in 2022. https://www.ftc.gov/news-events/news/press-releases/2023/02/'
    'new-ftc-data-show-consumers-reported-losing-nearly-88-billion-scams-2022',
    'Grover, P., Xu, J., Tittelfitz, J., Cheng, A., Li, Z., Zablocki, J., Liu, J., & Zhou, H. (2022). '
    'Fraud dataset benchmark and applications. arXiv. https://doi.org/10.48550/arXiv.2208.14417',
    'Harris, B. (2016). Sparkov Data Generation Tool [Source code]. GitHub. '
    'https://github.com/namebrandon/Sparkov_Data_Generation',
    'Kaggle/Shenoy, K. (2020). Credit Card Transactions Fraud Detection Dataset [Dataset]. '
    'https://www.kaggle.com/datasets/kartik2112/fraud-detection',
    'Lopez-Rojas, E. A., Elmir, A., & Axelsson, S. (2016). PaySim: A financial mobile money simulator for '
    'fraud detection. Proceedings of the 28th European Modeling and Simulation Symposium (EMSS 2016).',
    'OJK. (2026, Januari 19). Kejahatan siber melonjak 550 persen, OJK ingatkan pentingnya keamanan '
    'digital. ANTARA News. https://www.antaranews.com/berita/5363118',
    'Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., … Duchesnay, É. '
    '(2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, '
    '2825–2830.',
    'Saito, T., & Rehmsmeier, M. (2015). The precision-recall plot is more informative than the ROC plot '
    'when evaluating binary classifiers on imbalanced datasets. PLOS ONE, 10(3), e0118432.',
]
for r in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Mm(7)
    p.paragraph_format.first_line_indent = Mm(-7)
    _font(p.add_run(r), 10)

# ══════════════════════════════ LAMPIRAN ══════════════════════════════
heading('Lampiran A. Reproduksibilitas')
para('Seluruh hasil bersifat reproducible (random seed = 42) dan dihasilkan oleh dua notebook: '
     '01_eda.ipynb (EDA dan penyiapan data → data_clean.csv) dan 02_analysis.ipynb (pemodelan dan '
     'evaluasi). Gambar pada artikel ini diambil langsung dari keluaran kedua notebook tersebut.', indent=True)

# ── SAVE ──
out_path = os.path.join(OUT_DIR, 'Deteksi_Fraud_KNN_vs_LR_JUISI.docx')
doc.save(out_path)
print(f'Paper created: {out_path}')
print(f'Tables: {_tbl[0]} | Figures: {_fig[0]} | Equations: {_eq[0]}')
